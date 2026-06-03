from io import BytesIO
import html
import os
import re

import pandas as pd
import streamlit as st
from docx import Document
from docx.shared import RGBColor, Inches
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT


# ============================================================
# PAGE CONFIG
# ============================================================

st.set_page_config(
    page_title="Strategic District Field Guide",
    page_icon="📘",
    layout="wide",
    initial_sidebar_state="expanded",
)

BUILT_IN_WORKBOOK = "Texas Top 24 Research.xlsx"


# ============================================================
# STYLING
# ============================================================

CUSTOM_CSS = """
<style>
:root {
    --navy:#102a43;
    --blue:#1d4ed8;
    --blue-soft:#eff6ff;
    --blue-border:#bfdbfe;
    --green:#166534;
    --green-soft:#f0fdf4;
    --red:#991b1b;
    --red-soft:#fef2f2;
    --amber:#92400e;
    --amber-soft:#fffbeb;
    --slate:#334155;
    --slate-soft:#f8fafc;
    --page:#f3f6fa;
    --surface:#ffffff;
    --text:#0f172a;
    --muted:#475569;
    --border:#cbd5e1;
}

.stApp,
[data-testid="stAppViewContainer"],
[data-testid="stHeader"] {
    background: var(--page) !important;
    color: var(--text) !important;
}

html,
body,
p,
li,
span,
div,
label,
section,
article,
.stMarkdown,
.stMarkdown p,
.stMarkdown li,
.stMarkdown span,
.stMarkdown div {
    color: var(--text) !important;
}

.block-container {
    padding-top: 1.2rem;
    padding-bottom: 2.5rem;
    max-width: 1180px;
}

[data-testid="stSidebar"] {
    background: #ffffff !important;
    border-right: 1px solid var(--border) !important;
}

[data-testid="stSidebar"] * {
    color: var(--text) !important;
}

h1, h2, h3, h4, h5, h6 {
    color: var(--navy) !important;
}

.main-title {
    font-size: 2rem;
    font-weight: 850;
    color: var(--navy) !important;
    margin-bottom: .25rem;
    letter-spacing:-.02em;
}

.subtitle {
    color: var(--muted) !important;
    font-size: .98rem;
    margin-bottom: 1rem;
}

.section-title {
    color: var(--navy) !important;
    font-size: 1.25rem;
    font-weight: 800;
    margin: .6rem 0 .3rem;
}

.helper-text {
    color: var(--muted) !important;
    font-size: .9rem;
}

.stTabs [data-baseweb="tab-list"] {
    background: #ffffff !important;
    border: 1px solid var(--border) !important;
    border-radius: 14px !important;
    padding: .25rem !important;
}

.stTabs [data-baseweb="tab"] {
    color: var(--slate) !important;
    background: transparent !important;
    border-radius: 10px !important;
    font-weight: 700 !important;
}

.stTabs [aria-selected="true"] {
    background: var(--blue-soft) !important;
    color: var(--blue) !important;
}

[data-testid="stMetric"] {
    background: #ffffff !important;
    border: 1px solid var(--border) !important;
    border-radius: 16px !important;
    padding: .85rem !important;
    box-shadow: 0 4px 14px rgba(15, 23, 42, .05) !important;
}

[data-testid="stMetric"] * {
    color: var(--text) !important;
}

input,
textarea,
select,
[data-baseweb="input"],
[data-baseweb="select"],
[data-baseweb="textarea"] {
    background: #ffffff !important;
    color: var(--text) !important;
    border-color: var(--border) !important;
}

.stButton button,
.stDownloadButton button {
    background: var(--blue) !important;
    color: #ffffff !important;
    border: 1px solid var(--blue) !important;
    border-radius: 10px !important;
    font-weight: 750 !important;
}

.stButton button:hover,
.stDownloadButton button:hover {
    background: #1e40af !important;
    border-color: #1e40af !important;
    color: #ffffff !important;
}

[data-testid="stExpander"] {
    background: #ffffff !important;
    border: 1px solid var(--border) !important;
    border-radius: 14px !important;
}

[data-testid="stExpander"] * {
    color: var(--text) !important;
}

.card {
    border:1px solid var(--border);
    border-radius:20px;
    padding:1.05rem;
    margin-bottom:1rem;
    background:var(--surface);
    color: var(--text) !important;
    box-shadow:0 8px 24px rgba(15,23,42,.08);
}

.card,
.card * {
    color: var(--text) !important;
}

.card h3 {
    color:var(--navy) !important;
    margin-top:0;
    margin-bottom:.2rem;
    font-size:1.35rem;
    letter-spacing:-.01em;
}

.meta {
    color:var(--muted) !important;
    font-size:.88rem;
    margin-bottom:.75rem;
    line-height:1.45;
}

.summary-grid {
    display:grid;
    grid-template-columns: repeat(3, minmax(0,1fr));
    gap:.75rem;
    margin:.75rem 0 .6rem;
}

.summary-box {
    background:#ffffff;
    border:1px solid var(--border);
    border-radius:14px;
    padding:.75rem;
    box-shadow:0 3px 12px rgba(15,23,42,.04);
}

.summary-label {
    color:var(--muted) !important;
    font-size:.72rem;
    font-weight:850;
    text-transform:uppercase;
    letter-spacing:.05em;
    margin-bottom:.2rem;
}

.summary-value {
    color:var(--text) !important;
    font-size:.92rem;
    line-height:1.35;
}

.quickprep {
    background: #ffffff;
    border:1px solid var(--blue-border);
    border-left:5px solid var(--blue);
    border-radius:16px;
    padding:.9rem;
    margin:.7rem 0 .8rem;
    color: var(--text) !important;
    box-shadow: 0 4px 16px rgba(15, 23, 42, .05);
}

.quickprep,
.quickprep * {
    color: var(--text) !important;
}

.quickprep-title {
    color:var(--navy) !important;
    font-weight:850;
    font-size:.95rem;
    text-transform:uppercase;
    letter-spacing:.04em;
    margin-bottom:.25rem;
}

.lead {
    background:var(--blue-soft);
    border-left:4px solid var(--blue);
    border-radius:14px;
    padding:.75rem;
    margin:.6rem 0;
    line-height:1.45;
    color: var(--text) !important;
}

.badge,
.chip {
    display:inline-block;
    border-radius:999px;
    padding:.25rem .58rem;
    margin:.14rem .16rem .14rem 0;
    font-size:.76rem;
    font-weight:750;
    line-height:1.2;
    border:1px solid transparent;
}

.chip {
    background:var(--slate-soft);
    color:var(--slate) !important;
    border-color:#e2e8f0;
}

.tag-math {
    background:#dbeafe;
    color:#1e3a8a !important;
    border-color:#bfdbfe;
}

.tag-mtss {
    background:#ccfbf1;
    color:#134e4a !important;
    border-color:#99f6e4;
}

.tag-spedell {
    background:#ede9fe;
    color:#4c1d95 !important;
    border-color:#ddd6fe;
}

.tag-ccmr {
    background:#fef3c7;
    color:#78350f !important;
    border-color:#fde68a;
}

.tag-teacher {
    background:#f1f5f9;
    color:#1e293b !important;
    border-color:#e2e8f0;
}

.tag-hqim {
    background:#e0e7ff;
    color:#312e81 !important;
    border-color:#c7d2fe;
}

.tag-funding {
    background:#fef9c3;
    color:#713f12 !important;
    border-color:#fde68a;
}

.tag-relationship {
    background:#dcfce7;
    color:#14532d !important;
    border-color:#bbf7d0;
}

.tag-default {
    background:#eef2ff;
    color:#312e81 !important;
    border-color:#e0e7ff;
}

.priority {
    display:inline-block;
    border-radius:999px;
    padding:.28rem .62rem;
    font-size:.75rem;
    font-weight:850;
    vertical-align:middle;
}

.priority-very-high {
    background:var(--red-soft);
    color:var(--red) !important;
    border: 1px solid #fecaca;
}

.priority-high {
    background:var(--green-soft);
    color:var(--green) !important;
    border: 1px solid #bbf7d0;
}

.priority-medium-high {
    background:var(--amber-soft);
    color:var(--amber) !important;
    border: 1px solid #fde68a;
}

.priority-medium {
    background:#e0f2fe;
    color:#075985 !important;
    border: 1px solid #bae6fd;
}

.metric-note {
    color:var(--muted) !important;
    font-size:.84rem;
    margin-top:-.25rem;
}

.howto-box {
    background:#ffffff;
    color: var(--text) !important;
    border:1px solid var(--border);
    border-radius:18px;
    padding:1rem 1.1rem;
    margin:.75rem 0;
    box-shadow:0 4px 16px rgba(15,23,42,.05);
}

.howto-box,
.howto-box * {
    color: var(--text) !important;
}

hr {
    border:0;
    border-top:1px solid var(--border);
    margin:1.2rem 0;
}

@media (max-width: 800px) {
    .summary-grid {
        grid-template-columns: 1fr;
    }
}
</style>
"""
st.markdown(CUSTOM_CSS, unsafe_allow_html=True)


# ============================================================
# EXPECTED LARGER WORKBOOK SHEETS
# ============================================================

SHEET_STRATEGIC = "Strategic Indicators"
SHEET_SCORECARD = "Master Scorecard"
SHEET_BASIC = "Basic District Info"
SHEET_CONTACTS = "District Contacts"
SHEET_LEADERSHIP = "Leadership and Governance"
SHEET_CSI = "CSI"
SHEET_TSI = "TSI"
DISTRICT_COL = "District Name"


# ============================================================
# BASIC HELPERS
# ============================================================

def safe_html(value):
    return html.escape(str(value))


def normalize_text(value):
    if value is None:
        return ""
    try:
        if pd.isna(value):
            return ""
    except Exception:
        pass
    return str(value).strip()


def normalize_key(value):
    text = normalize_text(value).lower()
    text = text.replace("\n", " ").replace("\r", " ").replace("\xa0", " ")
    text = re.sub(r"\s+", " ", text).strip()
    return text


def district_key(value):
    return normalize_text(value).upper().strip()


def make_unique_columns(columns):
    seen = {}
    output = []
    for col in columns:
        base = normalize_text(col) or "Unnamed"
        if base not in seen:
            seen[base] = 0
            output.append(base)
        else:
            seen[base] += 1
            output.append(f"{base}.{seen[base]}")
    return output


def clean_columns(df):
    if df.empty:
        return df
    df = df.copy()
    df.columns = make_unique_columns(
        [str(c).strip().replace("\n", " ").replace("\r", " ") for c in df.columns]
    )
    drop_cols = [
        c for c in df.columns
        if normalize_key(c).startswith("methods:")
        or normalize_key(c).startswith("unnamed")
    ]
    if drop_cols:
        df = df.drop(columns=drop_cols)
    return df


def find_sheet_name(xls, desired_name):
    lookup = {normalize_key(sheet): sheet for sheet in xls.sheet_names}
    return lookup.get(normalize_key(desired_name))


def read_table_sheet(xls, desired_sheet_name):
    actual_sheet = find_sheet_name(xls, desired_sheet_name)
    if not actual_sheet:
        return pd.DataFrame()

    raw = pd.read_excel(xls, sheet_name=actual_sheet, header=None, engine="openpyxl")
    if raw.empty:
        return pd.DataFrame()

    header_row = None
    for idx, row in raw.iterrows():
        values = [normalize_text(v) for v in row.tolist()]
        if DISTRICT_COL in values:
            header_row = idx
            break

    if header_row is None:
        return pd.DataFrame()

    headers = [normalize_text(v) for v in raw.iloc[header_row].tolist()]
    data = raw.iloc[header_row + 1:].copy()
    data.columns = headers
    data = clean_columns(data)
    data = data.dropna(how="all")

    if DISTRICT_COL in data.columns:
        data = data[data[DISTRICT_COL].notna()]
        data = data[data[DISTRICT_COL].astype(str).str.strip() != ""]

    return data.reset_index(drop=True)


def find_col(df, exact_names=None, contains_all=None):
    if df.empty:
        return None
    exact_names = exact_names or []
    contains_all = contains_all or []
    normalized_map = {normalize_key(col): col for col in df.columns}

    for name in exact_names:
        key = normalize_key(name)
        if key in normalized_map:
            return normalized_map[key]

    if contains_all:
        for col in df.columns:
            col_key = normalize_key(col)
            if all(term.lower() in col_key for term in contains_all):
                return col

    return None


def get_value(row_or_dict, col_name, default=""):
    if row_or_dict is None:
        return default

    if isinstance(row_or_dict, dict):
        for key, value in row_or_dict.items():
            if normalize_key(key) == normalize_key(col_name):
                return value
        return default

    try:
        for key in row_or_dict.index:
            if normalize_key(key) == normalize_key(col_name):
                return row_or_dict.get(key, default)
    except Exception:
        return default

    return default


def format_number(value, decimals=2):
    if normalize_text(value) == "":
        return ""
    try:
        return f"{float(value):.{decimals}f}"
    except Exception:
        return normalize_text(value)


def format_enrollment(value):
    if normalize_text(value) == "":
        return ""
    try:
        return f"{int(float(value)):,}"
    except Exception:
        return normalize_text(value)


def format_percent(value, decimals=0):
    if normalize_text(value) == "":
        return ""
    try:
        number = float(value)
        pct = number * 100 if abs(number) <= 1 else number
        if decimals == 0:
            return f"{pct:.0f}%"
        return f"{pct:.{decimals}f}%"
    except Exception:
        return normalize_text(value)


def truncate_text(text, max_chars=190):
    text = normalize_text(text)
    if len(text) <= max_chars:
        return text
    cut = text[:max_chars].rsplit(" ", 1)[0]
    return cut + "..."


def format_contact_display(name="", title="", email="", position=""):
    name = normalize_text(name)
    title = normalize_text(title) or normalize_text(position)
    email = normalize_text(email)

    if name and title and email:
        return f"{name} — {title} | {email}"
    if name and title:
        return f"{name} — {title}"
    if name and email:
        return f"{name} | {email}"
    if title and email:
        return f"{title} | {email}"
    return name or title or email


def tag_class(tag):
    mapping = {
        "Math": "tag-math",
        "MTSS": "tag-mtss",
        "SPED/ELL": "tag-spedell",
        "CCMR": "tag-ccmr",
        "Teacher Capacity": "tag-teacher",
        "Curriculum / HQIM": "tag-hqim",
        "Funding / Grants": "tag-funding",
        "Existing Relationship": "tag-relationship",
    }
    return mapping.get(tag, "tag-default")


def badge_html(tag):
    return f'<span class="badge {tag_class(tag)}">{safe_html(tag)}</span>'


def chip_html(item):
    return f'<span class="chip">{safe_html(item)}</span>'


# ============================================================
# COLUMN DETECTION
# ============================================================

def detect_scorecard_columns(scorecard_df):
    return {
        "district": find_col(scorecard_df, exact_names=["District Name"]),
        "enrollment": find_col(scorecard_df, exact_names=["Enrollment"]),
        "strategic_score": find_col(
            scorecard_df,
            exact_names=["Strategic Weighted Score"],
            contains_all=["strategic", "weighted", "score"],
        ),
        "overall_score": find_col(
            scorecard_df,
            exact_names=[
                "Overall Weighted Score (Strategic+Contract+Relationship)",
                "Overall Weighted Score (Strategic+Relationship Weighted)",
                "Overall Weighted Score",
            ],
            contains_all=["overall", "weighted", "score"],
        ),
        "tier": find_col(scorecard_df, exact_names=["Tier"]),
        "existing_contracts": find_col(
            scorecard_df,
            exact_names=["Existing Contracts in Region (Yes/No)"],
        ),
        "existing_relationships": find_col(
            scorecard_df,
            exact_names=["Existing Relationships"],
        ),
    }


# ============================================================
# LOOKUPS AND ELIGIBILITY
# ============================================================

def is_substantive_strategic_row(row):
    fields_to_check = [
        "Strategic Plan Themes",
        "Math Improvement Mentioned",
        "Math Priority Strength",
        "Intervention Focus",
        "Intervention Focus Details",
        "Teacher Capacity/PD Focus",
        "Teacher Capacity Details",
        "Career Readiness Mentioned",
        "Career Readiness Details",
        "SPED/ELL Improvement Mentioned",
        "SPED/ELL Details",
        "MTSS/Tiered Support Mentioned",
        "MTSS Details",
        "Curriculum Review/Adoption Activity",
        "Curriculum Details",
        "Active Grants (Yes/No)",
        "Grants Details",
        "Sources",
        "Notes",
    ]
    return sum(1 for field in fields_to_check if normalize_text(get_value(row, field))) >= 2


def build_score_lookup(scorecard_df, score_cols):
    lookup = {}
    district_col = score_cols.get("district")
    if scorecard_df.empty or not district_col:
        return lookup
    for _, row in scorecard_df.iterrows():
        district = normalize_text(row.get(district_col, ""))
        if district:
            lookup[district_key(district)] = row.to_dict()
    return lookup


def build_basic_lookup(basic_df):
    lookup = {}
    if basic_df.empty or DISTRICT_COL not in basic_df.columns:
        return lookup
    for _, row in basic_df.iterrows():
        district = normalize_text(row.get(DISTRICT_COL, ""))
        if district:
            lookup[district_key(district)] = row.to_dict()
    return lookup


def build_csi_tsi_counts(csi_df, tsi_df):
    csi_counts = {}
    tsi_counts = {}
    if not csi_df.empty and DISTRICT_COL in csi_df.columns:
        csi_counts = (
            csi_df.groupby(csi_df[DISTRICT_COL].astype(str).str.upper().str.strip())
            .size()
            .to_dict()
        )
    if not tsi_df.empty and DISTRICT_COL in tsi_df.columns:
        tsi_counts = (
            tsi_df.groupby(tsi_df[DISTRICT_COL].astype(str).str.upper().str.strip())
            .size()
            .to_dict()
        )
    return csi_counts, tsi_counts


def determine_priority(tier, score):
    if tier == "Tier 1":
        return "Very High"
    if tier == "Tier 2":
        return "High"
    try:
        if float(score) >= 3.5:
            return "Medium-High"
    except Exception:
        pass
    return "Medium"


# ============================================================
# CONTACTS
# ============================================================

def build_contacts_lookup_from_contacts(contacts_df):
    lookup = {}
    if contacts_df.empty or DISTRICT_COL not in contacts_df.columns:
        return lookup

    role_priority = [
        "SUPERINTENDENT",
        "CHIEF ACADEMIC OFFICER",
        "ASSISTANT SUPERINTENDENT",
        "CURRICULUM DIRECTOR",
        "MATH COORDINATOR",
        "CTE COLLEGE CAREER READINESS DIRECTOR",
        "SPECIAL EDUCATION DIRECTOR",
        "ESL ELL COORDINATOR",
        "DIRECTOR ASSESSMENT DATA",
        "PROFESSIONAL LEARNING DIRECTOR",
        "ELA LITERACY COORDINATOR",
    ]

    df = contacts_df.copy()
    if "Position" in df.columns:
        df["_priority"] = (
            df["Position"]
            .astype(str)
            .str.upper()
            .apply(lambda x: role_priority.index(x) if x in role_priority else 99)
        )
        df = df.sort_values([DISTRICT_COL, "_priority"])

    for _, row in df.iterrows():
        district = normalize_text(row.get(DISTRICT_COL, ""))
        if not district:
            continue

        first = normalize_text(row.get("First Name", ""))
        last = normalize_text(row.get("Last Name", ""))
        title = normalize_text(row.get("Title", ""))
        position = normalize_text(row.get("Position", ""))
        email = normalize_text(row.get("Email", ""))

        name = f"{first} {last}".strip()
        contact = format_contact_display(
            name=name,
            title=title,
            email=email,
            position=position,
        )

        if contact:
            lookup.setdefault(district_key(district), [])
            if contact not in lookup[district_key(district)]:
                lookup[district_key(district)].append(contact)

    return lookup


def build_contacts_lookup_from_leadership(leadership_df):
    lookup = {}
    if leadership_df.empty or DISTRICT_COL not in leadership_df.columns:
        return lookup

    for _, row in leadership_df.iterrows():
        district = normalize_text(row.get(DISTRICT_COL, ""))
        if not district:
            continue

        contacts = []

        superintendent = normalize_text(get_value(row, "Superintendent"))
        superintendent_email = normalize_text(get_value(row, "Email"))

        curriculum_lead = normalize_text(get_value(row, "Curriculum Lead"))
        curriculum_title = normalize_text(get_value(row, "Title")) or "Curriculum / Academic Lead"
        curriculum_email = normalize_text(get_value(row, "Email.1"))

        cte_lead = normalize_text(get_value(row, "CTE Lead"))
        cte_title = normalize_text(get_value(row, "Title.1")) or "CTE / Career Readiness Lead"
        cte_email = normalize_text(get_value(row, "Email.2"))

        math_lead = normalize_text(get_value(row, "Math Lead"))
        math_title = normalize_text(get_value(row, "Title.2")) or "Math Lead"
        math_email = normalize_text(get_value(row, "Email.3"))

        for contact in [
            format_contact_display(superintendent, "Superintendent", superintendent_email),
            format_contact_display(curriculum_lead, curriculum_title, curriculum_email),
            format_contact_display(cte_lead, cte_title, cte_email),
            format_contact_display(math_lead, math_title, math_email),
        ]:
            if contact:
                contact = contact.replace("[", "").replace("]", "")
                contact = contact.replace("(mailto:", " | ").replace(")", "")
                contacts.append(contact)

        lookup[district_key(district)] = contacts[:6]

    return lookup


def get_contacts(district_name, contacts_lookup, leadership_lookup):
    key = district_key(district_name)
    if key in contacts_lookup and contacts_lookup[key]:
        return contacts_lookup[key][:6]
    return leadership_lookup.get(key, [])[:6]


# ============================================================
# STRATEGIC CARD LANGUAGE
# ============================================================

def infer_tags(strategic_row, score_row):
    tags = []
    combined_text = " ".join(
        normalize_text(v)
        for v in strategic_row.to_dict().values()
    ).lower()

    if normalize_text(get_value(strategic_row, "Math Improvement Mentioned")) or normalize_text(get_value(strategic_row, "Math Priority Strength")):
        tags.append("Math")
    if normalize_text(get_value(strategic_row, "Intervention Focus")) or normalize_text(get_value(strategic_row, "MTSS/Tiered Support Mentioned")):
        tags.append("MTSS")
    if (
        "sped" in combined_text
        or "special education" in combined_text
        or "english learner" in combined_text
        or "emergent bilingual" in combined_text
        or "ell" in combined_text
        or "multilingual" in combined_text
    ):
        tags.append("SPED/ELL")
    if normalize_text(get_value(strategic_row, "Career Readiness Mentioned")) or normalize_text(get_value(strategic_row, "Career Readiness Details")):
        tags.append("CCMR")
    if normalize_text(get_value(strategic_row, "Teacher Capacity/PD Focus")) or normalize_text(get_value(strategic_row, "Teacher Capacity Details")):
        tags.append("Teacher Capacity")
    if normalize_text(get_value(strategic_row, "Curriculum Review/Adoption Activity")) or normalize_text(get_value(strategic_row, "Curriculum Details")):
        tags.append("Curriculum / HQIM")
    if normalize_text(get_value(strategic_row, "Active Grants (Yes/No)")) or normalize_text(get_value(strategic_row, "Grants Details")):
        tags.append("Funding / Grants")
    if normalize_text(get_value(score_row, "Existing Relationships")).lower() == "yes":
        tags.append("Existing Relationship")

    if not tags:
        tags.append("Strategic Review")

    unique = []
    for tag in tags:
        if tag not in unique:
            unique.append(tag)
    return unique


def build_alignment(tags):
    alignment = []
    if "Math" in tags:
        alignment.append("Elevation Station Math Games — K–8 math practice, fluency, reinforcement, and engagement.")
        alignment.append("Elevation intervention curriculum — targeted K–5 support where Tier 2/Tier 3 math or foundational gaps are present.")
    if "MTSS" in tags:
        alignment.append("PCG MTSS consulting — system design, campus implementation, intervention fidelity, and data-to-action routines.")
    if "SPED/ELL" in tags:
        alignment.append("PCG SPED / multilingual learner support — inclusive practice, service delivery, access to grade-level instruction, and subgroup progress monitoring.")
    if "CCMR" in tags:
        alignment.append("RISE Career & Math Mini Lessons — grades 6–9 career-connected math, pathway awareness, and applied readiness.")
    if "Curriculum / HQIM" in tags:
        alignment.append("PCG curriculum/HQIM implementation support — adoption fidelity, coaching, PLC routines, and change management.")
    if "Teacher Capacity" in tags:
        alignment.append("PCG professional learning and instructional implementation support — teacher capacity, coaching, data use, and scalable instructional routines.")
    if "Funding / Grants" in tags:
        alignment.append("Funding alignment support — connect implementation supports to existing grant, Title, or strategic funding streams where appropriate.")
    if not alignment:
        alignment.append("Discovery needed — validate strategic needs, stakeholder priorities, and fit before positioning specific resources.")
    return alignment


def build_lead_with(card):
    tags = card.get("tags", [])
    if "Existing Relationship" in tags and card.get("tier") == "Tier 1":
        relationship_phrase = "given the existing relationship and strong strategic fit"
    elif "Existing Relationship" in tags:
        relationship_phrase = "building from the existing relationship"
    else:
        relationship_phrase = "as an initial consultative entry point"
    top_tags = [tag for tag in tags if tag not in ["Existing Relationship", "Funding / Grants"]]
    top_tags_text = ", ".join(top_tags[:3]).lower() or "strategic implementation support"
    return f"Lead with {top_tags_text} {relationship_phrase}."


def build_refined_questions(card):
    tags = card.get("tags", [])
    questions = []

    if "Math" in tags:
        questions.append("How are campuses currently using assessment data to decide which students need additional math support?")
    else:
        questions.append("How are campuses currently translating district priorities into daily instructional routines?")

    questions.append("Where does implementation tend to vary most across campuses, grade levels, or student groups?")

    if "MTSS" in tags:
        questions.append("What tends to slow down the response after students are identified for additional support?")

    if "SPED/ELL" in tags:
        questions.append("Where do students with disabilities or multilingual learners need more consistent access to grade-level expectations?")

    if "Teacher Capacity" in tags or "Curriculum / HQIM" in tags:
        questions.append("What makes it easier or harder for teachers to use new materials or supports consistently after initial training?")

    if "CCMR" in tags:
        questions.append("Where do students begin connecting academic skills to future pathways and career readiness expectations?")

    questions.append("What evidence would tell district and campus leaders that a support model is working well enough to expand?")
    questions.append("If a small pilot were considered, what would make the pilot credible to teachers and principals?")

    unique = []
    for q in questions:
        if q not in unique:
            unique.append(q)
    return unique[:6]


def build_listen_for(tags):
    tag_map = {
        "Math": ["math growth", "early numeracy", "Algebra readiness", "STAAR math", "student practice"],
        "MTSS": ["intervention fidelity", "Tier 2", "Tier 3", "progress monitoring", "campus variation"],
        "SPED/ELL": ["access to grade-level instruction", "service delivery", "emergent bilingual students", "students with disabilities", "subgroup gaps"],
        "CCMR": ["pathways", "industry-based certifications", "TSIA2", "dual credit", "career awareness"],
        "Teacher Capacity": ["teacher burden", "coaching", "professional learning", "PLC routines", "instructional consistency"],
        "Curriculum / HQIM": ["adoption fidelity", "HQIM", "curriculum implementation", "Eureka", "Bluebonnet", "instructional materials"],
        "Funding / Grants": ["Title funding", "grant alignment", "LASSO", "federal funds", "implementation funding"],
        "Existing Relationship": ["existing relationship", "current contract", "expansion", "trusted partner"],
    }
    listen_for = []
    for tag in tags:
        listen_for.extend(tag_map.get(tag, []))
    listen_for.extend(["data cycles", "implementation barriers", "capacity constraints"])
    unique = []
    for item in listen_for:
        if item not in unique:
            unique.append(item)
    return unique


def build_avoid(tags):
    avoid = [
        "Do not lead with a product pitch.",
        "Avoid positioning support as a replacement for the district’s current strategy or adopted curriculum.",
    ]
    if "Curriculum / HQIM" in tags:
        avoid.append("Frame support around implementation and teacher usability rather than another curriculum.")
    if "SPED/ELL" in tags:
        avoid.append("Avoid treating subgroup performance as a side issue; connect it to access and implementation.")
    if "Existing Relationship" in tags:
        avoid.append("Build from existing relationship context before introducing a new idea.")
    return avoid


# ============================================================
# SUMMARY BUILDERS
# ============================================================

def get_priority_need_opportunity(tags):
    items = []
    if "Math" in tags:
        items.append("Priority: math growth. Need: consistent practice and progress monitoring. Opportunity: Elevation Station and targeted math intervention.")
    if "MTSS" in tags:
        items.append("Priority: intervention. Need: clear Tier 2/Tier 3 routines. Opportunity: PCG MTSS support and intervention design.")
    if "SPED/ELL" in tags:
        items.append("Priority: subgroup access. Need: grade-level expectations with usable scaffolds. Opportunity: PCG SPED/ML support and aligned practice tools.")
    if "CCMR" in tags:
        items.append("Priority: readiness pathways. Need: earlier career-connected relevance. Opportunity: RISE Career & Math Mini Lessons.")
    if "Curriculum / HQIM" in tags or "Teacher Capacity" in tags:
        items.append("Priority: implementation quality. Need: teacher-ready routines after training. Opportunity: PCG implementation support and Emerald practice tools.")
    if not items:
        items.append("Priority: validate fit. Need: clarify district pain points and implementation constraints. Opportunity: discovery conversation.")
    return items[:3]


def build_relationship_insights(card):
    tags = card.get("tags", [])
    insights = []
    if "Existing Relationship" in tags:
        insights.append("Build from existing relationship credibility before introducing new support.")
    else:
        insights.append("Start with district priorities and implementation realities before discussing resources.")
    if "Curriculum / HQIM" in tags:
        insights.append("Avoid positioning as a competing curriculum; focus on adoption support and classroom usability.")
    elif "Teacher Capacity" in tags:
        insights.append("Emphasize reducing teacher burden and strengthening routines already expected by the district.")
    else:
        insights.append("Use consultative language around practical implementation and measurable proof points.")
    insights.append("Best next step is a narrow proof point tied to a campus, grade band, or student group need.")
    return insights


def build_compact_summary(card):
    tags = card.get("tags", [])
    signals = card.get("signals", [])
    top_tags = [t for t in tags if t not in ["Existing Relationship", "Funding / Grants"]]

    why = f"High-interest fit around {', '.join(top_tags[:3]).lower()}." if top_tags else "Strategic fit should be validated through discovery."
    if card.get("priority") in ["Very High", "High"]:
        why = f"{card.get('priority')} priority account with signals around {', '.join(top_tags[:3]).lower()}." if top_tags else f"{card.get('priority')} priority account."

    entry = get_priority_need_opportunity(tags)[0]

    barrier = "Likely barrier: campus variation, teacher capacity, or implementation consistency."
    if "Curriculum / HQIM" in tags:
        barrier = "Likely barrier: turning curriculum or HQIM expectations into consistent classroom routines."
    elif "MTSS" in tags:
        barrier = "Likely barrier: making intervention routines consistent across campuses."
    elif "SPED/ELL" in tags:
        barrier = "Likely barrier: maintaining grade-level access while differentiating support."

    return {
        "why": why,
        "entry": entry,
        "barrier": barrier,
        "top_signals": [truncate_text(s, 210) for s in signals[:3]],
        "positioning": get_priority_need_opportunity(tags),
        "relationship": build_relationship_insights(card),
    }


# ============================================================
# CARD BUILDING
# ============================================================

def build_card(strategic_row, score_row, basic_lookup, contacts_lookup, leadership_lookup, csi_counts, tsi_counts, score_cols):
    district_name = normalize_text(get_value(strategic_row, DISTRICT_COL))
    key = district_key(district_name)

    raw_score = score_row.get(score_cols["overall_score"], "")
    score = format_number(raw_score, 2)
    strategic_score = format_number(score_row.get(score_cols["strategic_score"], ""), 2) if score_cols.get("strategic_score") else ""
    tier = normalize_text(score_row.get(score_cols["tier"], ""))
    enrollment = format_enrollment(score_row.get(score_cols["enrollment"], "")) if score_cols.get("enrollment") else ""

    tags = infer_tags(strategic_row, score_row)
    priority = determine_priority(tier, raw_score)
    basic = basic_lookup.get(key, {})

    signals = []

    themes = normalize_text(get_value(strategic_row, "Strategic Plan Themes"))
    math_strength = normalize_text(get_value(strategic_row, "Math Priority Strength"))
    intervention_details = normalize_text(get_value(strategic_row, "Intervention Focus Details"))
    sped_ell_details = normalize_text(get_value(strategic_row, "SPED/ELL Details"))
    teacher_details = normalize_text(get_value(strategic_row, "Teacher Capacity Details"))
    career_details = normalize_text(get_value(strategic_row, "Career Readiness Details"))
    mtss_details = normalize_text(get_value(strategic_row, "MTSS Details"))
    curriculum_details = normalize_text(get_value(strategic_row, "Curriculum Details"))
    grants_details = normalize_text(get_value(strategic_row, "Grants Details"))

    if themes:
        signals.append(f"Strategic themes: {themes}")
    if math_strength:
        signals.append(f"Math signal: {math_strength}")

    csi = normalize_text(get_value(basic, "CSI Schools")) or str(csi_counts.get(key, ""))
    tsi = normalize_text(get_value(basic, "TSI Schools")) or str(tsi_counts.get(key, ""))
    if csi or tsi:
        signals.append(f"Accountability pressure: {csi or '0'} CSI schools and {tsi or '0'} TSI schools.")

    context_parts = []
    context_fields = [
        ("schools", "Number of Schools", "text"),
        ("grade span", "Grade Span Served", "text"),
        ("setting", "Urban/Suburban/Rural", "text"),
        ("economically disadvantaged", "% Economically Disadvantaged", "percent"),
        ("English learner", "% English Learner", "percent"),
        ("special education", "% Special Education", "percent"),
        ("major student group", "Major Student Groups", "text"),
        ("student growth trend", "Student Growth Trend", "percent1"),
    ]

    for label, col, kind in context_fields:
        raw_value = get_value(basic, col)
        value = normalize_text(raw_value)
        if value:
            if kind == "percent":
                value = format_percent(raw_value, 0)
            elif kind == "percent1":
                value = format_percent(raw_value, 1)
            context_parts.append(f"{label}: {value}")

    if context_parts:
        signals.append("District context: " + "; ".join(context_parts) + ".")

    if intervention_details:
        signals.append(f"Intervention signal: {intervention_details}")
    if mtss_details:
        signals.append(f"MTSS signal: {mtss_details}")
    if sped_ell_details:
        signals.append(f"SPED/ELL signal: {sped_ell_details}")
    if teacher_details:
        signals.append(f"Teacher capacity signal: {teacher_details}")
    if career_details:
        signals.append(f"CCMR / career readiness signal: {career_details}")
    if curriculum_details:
        signals.append(f"Curriculum / implementation signal: {curriculum_details}")
    if grants_details:
        signals.append(f"Funding / grants signal: {grants_details}")

    existing_relationships = normalize_text(get_value(score_row, "Existing Relationships"))
    existing_contracts = normalize_text(get_value(score_row, "Existing Contracts in Region (Yes/No)"))

    if existing_contracts and existing_relationships:
        signals.append(f"Relationship context: existing contracts in region = {existing_contracts}; existing relationships = {existing_relationships}.")
    elif existing_relationships:
        signals.append(f"Relationship context: existing relationships = {existing_relationships}.")
    elif existing_contracts:
        signals.append(f"Relationship context: existing contracts in region = {existing_contracts}.")

    card = {
        "name": district_name,
        "tier": tier,
        "score": score,
        "strategic_score": strategic_score,
        "enrollment": enrollment,
        "priority": priority,
        "tags": tags,
        "contacts": get_contacts(district_name, contacts_lookup, leadership_lookup),
        "signals": signals,
        "alignment": build_alignment(tags),
        "listen": build_listen_for(tags),
        "avoid": build_avoid(tags),
    }

    card["lead"] = build_lead_with(card)
    card["questions"] = build_refined_questions(card)
    card["compact"] = build_compact_summary(card)

    return card


# ============================================================
# LOAD WORKBOOK
# ============================================================

def get_workbook_source(uploaded_file):
    if uploaded_file is not None:
        return uploaded_file, "Uploaded workbook override"
    if os.path.exists(BUILT_IN_WORKBOOK):
        return BUILT_IN_WORKBOOK, "Built-in workbook"
    return None, "No workbook found"


def load_cards_from_workbook(workbook_source):
    xls = pd.ExcelFile(workbook_source, engine="openpyxl")

    strategic_df = read_table_sheet(xls, SHEET_STRATEGIC)
    scorecard_df = read_table_sheet(xls, SHEET_SCORECARD)
    basic_df = read_table_sheet(xls, SHEET_BASIC)
    contacts_df = read_table_sheet(xls, SHEET_CONTACTS)
    leadership_df = read_table_sheet(xls, SHEET_LEADERSHIP)
    csi_df = read_table_sheet(xls, SHEET_CSI)
    tsi_df = read_table_sheet(xls, SHEET_TSI)

    if strategic_df.empty or scorecard_df.empty:
        return [], xls, strategic_df, scorecard_df, basic_df, contacts_df, leadership_df, csi_df, tsi_df, pd.DataFrame(), {}

    score_cols = detect_scorecard_columns(scorecard_df)
    score_lookup = build_score_lookup(scorecard_df, score_cols)
    basic_lookup = build_basic_lookup(basic_df)
    contacts_lookup = build_contacts_lookup_from_contacts(contacts_df)
    leadership_lookup = build_contacts_lookup_from_leadership(leadership_df)
    csi_counts, tsi_counts = build_csi_tsi_counts(csi_df, tsi_df)

    cards = []
    audit_rows = []

    for _, strategic_row in strategic_df.iterrows():
        district_name = normalize_text(get_value(strategic_row, DISTRICT_COL))
        if not district_name:
            continue

        key = district_key(district_name)
        substantive = is_substantive_strategic_row(strategic_row)
        score_row = score_lookup.get(key, {})
        matched = bool(score_row)
        tier = normalize_text(score_row.get(score_cols.get("tier"), "")) if matched and score_cols.get("tier") else ""
        overall = score_row.get(score_cols.get("overall_score"), "") if matched and score_cols.get("overall_score") else ""
        eligible = substantive and matched and bool(tier) and bool(normalize_text(overall))

        audit_rows.append({
            "District": district_name,
            "Substantive Strategic Indicators": substantive,
            "Matched Master Scorecard": matched,
            "Detected Tier Column": score_cols.get("tier"),
            "Tier": tier,
            "Detected Overall Score Column": score_cols.get("overall_score"),
            "Overall Score": normalize_text(overall),
            "Eligible": eligible,
        })

        if eligible:
            cards.append(
                build_card(
                    strategic_row,
                    score_row,
                    basic_lookup,
                    contacts_lookup,
                    leadership_lookup,
                    csi_counts,
                    tsi_counts,
                    score_cols,
                )
            )

    audit_df = pd.DataFrame(audit_rows)
    return cards, xls, strategic_df, scorecard_df, basic_df, contacts_df, leadership_df, csi_df, tsi_df, audit_df, score_cols


# ============================================================
# DEBUG
# ============================================================

def show_workbook_debug(xls, strategic_df, scorecard_df, basic_df, contacts_df, leadership_df, csi_df, tsi_df, audit_df, score_cols, data_source_label):
    st.markdown('<div class="section-title">Workbook Diagnostics</div>', unsafe_allow_html=True)
    st.markdown(
        f'<div class="helper-text">Current data source: <strong>{safe_html(data_source_label)}</strong></div>',
        unsafe_allow_html=True,
    )
    st.markdown("### Sheets found")
    st.write(xls.sheet_names)
    st.markdown("### Detected scorecard columns")
    st.write(score_cols)

    for label, df in [
        ("Strategic Indicators", strategic_df),
        ("Master Scorecard", scorecard_df),
        ("Basic District Info", basic_df),
        ("District Contacts", contacts_df),
        ("Leadership and Governance", leadership_df),
        ("CSI", csi_df),
        ("TSI", tsi_df),
    ]:
        with st.expander(label, expanded=False):
            st.write(f"Rows: {len(df)}")
            st.write("Columns:")
            st.write(list(df.columns))
            if DISTRICT_COL in df.columns:
                st.write("First district names found:")
                st.write(df[DISTRICT_COL].dropna().astype(str).head(15).tolist())

    st.markdown("### Eligibility audit")
    st.dataframe(audit_df, use_container_width=True)


# ============================================================
# SEARCH / FILTER
# ============================================================

def search_blob(card):
    values = [
        card.get("name", ""),
        card.get("tier", ""),
        str(card.get("score", "")),
        str(card.get("strategic_score", "")),
        card.get("enrollment", ""),
        card.get("priority", ""),
        card.get("lead", ""),
    ]
    for key in ["tags", "signals", "alignment", "contacts", "questions", "listen", "avoid"]:
        values.extend(card.get(key, []))
    return " ".join(map(str, values)).lower()


def filter_cards(cards, query, selected_tiers, selected_tags, shortlist, show_shortlist_only):
    query = query.strip().lower()
    filtered = []

    for card in cards:
        if selected_tiers and card.get("tier") not in selected_tiers:
            continue
        if selected_tags and not any(tag in card.get("tags", []) for tag in selected_tags):
            continue
        if show_shortlist_only and card.get("name") not in shortlist:
            continue
        if query and query not in search_blob(card):
            continue
        filtered.append(card)

    return filtered


# ============================================================
# RENDERING
# ============================================================

def render_summary_grid(card):
    compact = card.get("compact", {})
    st.markdown(f"""
    <div class="summary-grid">
        <div class="summary-box">
            <div class="summary-label">Why it matters</div>
            <div class="summary-value">{safe_html(compact.get('why', ''))}</div>
        </div>
        <div class="summary-box">
            <div class="summary-label">Best entry point</div>
            <div class="summary-value">{safe_html(compact.get('entry', ''))}</div>
        </div>
        <div class="summary-box">
            <div class="summary-label">Likely barrier</div>
            <div class="summary-value">{safe_html(compact.get('barrier', ''))}</div>
        </div>
    </div>
    """, unsafe_allow_html=True)


def render_quick_prep(card):
    top_listen = card.get("listen", [])[:8]
    best_question = card.get("questions", [""])[0] if card.get("questions") else ""
    avoid = card.get("avoid", [""])[0] if card.get("avoid") else ""
    chips = "".join(chip_html(item) for item in top_listen)
    st.markdown(f"""
    <div class="quickprep">
        <div class="quickprep-title">Quick Prep</div>
        <strong>Best opening question:</strong> {safe_html(best_question)}<br>
        <strong>Listen for:</strong><br>{chips}<br>
        <strong>Avoid:</strong> {safe_html(avoid)}
    </div>
    """, unsafe_allow_html=True)


def render_card(card, view_mode="Quick Brief"):
    priority_class = card.get("priority", "Medium").lower().replace(" ", "-")
    badges = "".join(badge_html(tag) for tag in card.get("tags", []))
    compact = card.get("compact", {})

    st.markdown(f"""
        <div class="card">
            <h3>{safe_html(card['name'])} <span class="priority priority-{priority_class}">{safe_html(card.get('priority', ''))}</span></h3>
            <div class="meta">
                {safe_html(card.get('tier', ''))}
                | Overall Score {safe_html(card.get('score', ''))}
                | Strategic Score {safe_html(card.get('strategic_score', ''))}
                | Enrollment {safe_html(card.get('enrollment', ''))}
            </div>
            <div>{badges}</div>
        </div>
    """, unsafe_allow_html=True)

    render_summary_grid(card)
    render_quick_prep(card)

    st.markdown("**Top Signals**")
    for item in compact.get("top_signals", [])[:3]:
        st.markdown(f"- {item}")

    st.markdown("**Best Fit**")
    for item in compact.get("positioning", [])[:3]:
        st.markdown(f"- {item}")

    st.markdown("**Discovery Questions**")
    for item in card.get("questions", [])[:4]:
        st.markdown(f"- {item}")

    with st.expander("Key Contacts", expanded=False):
        contacts = card.get("contacts", [])
        if contacts:
            for item in contacts[:6]:
                st.markdown(f"- {item}")
        else:
            st.markdown("_No contacts were available for this district._")

    if view_mode == "Full Detail":
        with st.expander("Full Strategic Detail", expanded=False):
            st.markdown("**Strategic Signals**")
            for item in card.get("signals", []):
                st.markdown(f"- {item}")

            st.markdown("**PCG / Emerald Alignment**")
            for item in card.get("alignment", []):
                st.markdown(f"- {item}")

            st.markdown("**Listen For**")
            st.markdown("".join(chip_html(item) for item in card.get("listen", [])), unsafe_allow_html=True)

            st.markdown("**Avoid**")
            for item in card.get("avoid", []):
                st.markdown(f"- {item}")

    quick_prep = (
        f"{card['name']} Quick Prep\n\n"
        f"Why it matters: {compact.get('why', '')}\n"
        f"Best entry point: {compact.get('entry', '')}\n"
        f"Likely barrier: {compact.get('barrier', '')}\n\n"
        f"Best opening question: {(card.get('questions') or [''])[0]}\n\n"
        f"Listen for: {', '.join(card.get('listen', [])[:8])}"
    )

    with st.expander("Copy Quick Prep", expanded=False):
        st.text_area("Quick prep copy", quick_prep, height=150, key=f"quick_{card['name']}")

    st.download_button(
        "Download this district brief",
        data=build_docx([card]),
        file_name=f"{card['name'].replace(' ', '_')}_Brief.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        key=f"download_{card['name']}",
    )


# ============================================================
# WORD EXPORTS
# ============================================================

def add_bullets(cell, items, limit=None):
    use_items = items[:limit] if limit else items
    if not use_items:
        cell.text = ""
        return
    cell.text = ""
    for item in use_items:
        p = cell.add_paragraph(style=None)
        p.text = f"• {item}"


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def build_docx(cards):
    doc = Document()
    title = doc.add_heading("Strategic District Field Guide", 0)
    for run in title.runs:
        run.font.color.rgb = RGBColor(23, 54, 93)

    doc.add_paragraph(
        "Mobile conversation cards for conference prep: lead-with angle, strategic signals, contacts, PCG/Emerald alignment, and discovery questions."
    )

    for card in cards:
        compact = card.get("compact", {})
        doc.add_heading(card["name"], level=1)
        doc.add_paragraph(
            f"{card.get('tier', '')} | Overall Score {card.get('score', '')} | Strategic Score {card.get('strategic_score', '')} | Enrollment {card.get('enrollment', '')} | Priority {card.get('priority', '')}"
        )

        doc.add_heading("Quick Read", level=2)
        for item in [compact.get("why"), compact.get("entry"), compact.get("barrier")]:
            if item:
                doc.add_paragraph(item, style="List Bullet")

        for section_title, key in [
            ("Top Strategic Signals", "signals"),
            ("Best-Fit PCG / Emerald Alignment", "alignment"),
            ("Key Contacts", "contacts"),
            ("Discovery Questions", "questions"),
            ("Listen For", "listen"),
            ("Avoid", "avoid"),
        ]:
            doc.add_heading(section_title, level=2)
            for item in card.get(key, []):
                doc.add_paragraph(item, style="List Bullet")

    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio


def build_matrix_docx(cards):
    doc = Document()
    section = doc.sections[0]
    section.left_margin = Inches(0.35)
    section.right_margin = Inches(0.35)
    section.top_margin = Inches(0.4)
    section.bottom_margin = Inches(0.4)

    title = doc.add_heading("Texas District Strategic Positioning Matrix", 0)
    for run in title.runs:
        run.font.color.rgb = RGBColor(31, 78, 121)

    doc.add_paragraph("Scope: Currently filtered districts from the Strategic District Field Guide.")
    doc.add_paragraph("Designed for deeper review after quick conference scanning in the app.")

    table = doc.add_table(rows=1, cols=6)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    headers = [
        "District / Score / Key Contacts",
        "District Overview",
        "Strategic Signal Analysis",
        "Positioning Matrix",
        "Discovery Questions",
        "Relationship & Engagement Insights",
    ]

    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        shade_cell(hdr_cells[i], "1F4E79")
        for p in hdr_cells[i].paragraphs:
            for run in p.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.bold = True

    for card in cards:
        compact = card.get("compact", {})
        row_cells = table.add_row().cells

        for cell in row_cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP

        contacts = card.get("contacts", [])[:5]
        row_cells[0].text = (
            f"{card.get('name','')}\n"
            f"{card.get('tier','')} | Overall {card.get('score','')}\n"
            f"Enrollment: {card.get('enrollment','')}\n"
            f"Contacts: " + "; ".join(contacts)
        )

        row_cells[1].text = ""
        add_bullets(
            row_cells[1],
            [
                compact.get("why", ""),
                compact.get("entry", ""),
                compact.get("barrier", ""),
            ],
        )

        row_cells[2].text = ""
        add_bullets(row_cells[2], compact.get("top_signals", []), limit=3)

        row_cells[3].text = ""
        add_bullets(row_cells[3], compact.get("positioning", []), limit=3)

        row_cells[4].text = ""
        add_bullets(row_cells[4], card.get("questions", []), limit=6)

        row_cells[5].text = ""
        add_bullets(row_cells[5], compact.get("relationship", []), limit=3)

    doc.add_heading("PCG / Emerald Alignment Lens", level=1)

    lens = doc.add_table(rows=1, cols=3)
    lens.style = "Table Grid"

    lens_headers = [
        "Strategic Need",
        "Consultative Alignment",
        "Representative Resource / Service Fit",
    ]

    for i, h in enumerate(lens_headers):
        lens.rows[0].cells[i].text = h
        shade_cell(lens.rows[0].cells[i], "1F4E79")
        for p in lens.rows[0].cells[i].paragraphs:
            for run in p.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.bold = True

    lens_rows = [
        (
            "Math proficiency / early numeracy / Algebra readiness",
            "Implementation support, progress monitoring routines, supplemental practice, teacher-facing supports",
            "Elevation Station Math Games; Elevation intervention curriculum; PCG instructional implementation support",
        ),
        (
            "MTSS / Tier 2 / Tier 3 intervention",
            "System design, campus implementation, intervention fidelity, data-to-action routines",
            "PCG MTSS consulting; Emerald intervention supports",
        ),
        (
            "SPED / ELL / subgroup performance",
            "Inclusive practice, service delivery, compliance-to-instruction alignment, differentiated supports",
            "PCG SPED and multilingual learner services; scaffolded Emerald resources",
        ),
        (
            "CCMR / career-connected learning",
            "Middle-grade pathway exposure, career relevance, readiness pipeline alignment",
            "RISE Career & Math Mini Lessons; PCG CCMR strategy support",
        ),
        (
            "Teacher capacity / HQIM or curriculum adoption",
            "Professional learning, coaching, PLC routines, adoption fidelity, change management",
            "PCG implementation and professional learning services; Emerald teacher-ready practice tools",
        ),
    ]

    for a, b, c in lens_rows:
        cells = lens.add_row().cells
        cells[0].text = a
        cells[1].text = b
        cells[2].text = c

    doc.add_paragraph(
        "Note: This matrix is designed for strategy review and conference/meeting preparation. Validate against current district conversations before final pursuit decisions."
    )

    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio


# ============================================================
# HOW-TO GUIDE
# ============================================================

def render_how_to(data_source_label):
    st.markdown('<div class="section-title">How-To Guide</div>', unsafe_allow_html=True)

    st.markdown(f"""
    <div class="howto-box">
    <strong>Current data source:</strong> {safe_html(data_source_label)}<br>
    The app opens with the built-in workbook from GitHub when available. Uploading a workbook in the sidebar temporarily overrides the built-in data for that session.
    </div>
    """, unsafe_allow_html=True)

    st.markdown("""
    <div class="howto-box">
    <strong>60-second conference workflow</strong>
    <ol>
      <li>Search the district name.</li>
      <li>Read Why It Matters, Best Entry Point, and Likely Barrier.</li>
      <li>Use the Quick Prep question to start discovery.</li>
      <li>Open Key Contacts for names, titles, and emails.</li>
      <li>Use the Full Matrix download for deeper post-meeting planning.</li>
    </ol>
    </div>
    """, unsafe_allow_html=True)

    st.markdown("### Useful searches")
    st.code("Dallas\nFort Worth\nSPED\nMTSS\nEureka\nBluebonnet\ncareer\nteacher burden", language="text")

    st.markdown("### Tag legend")
    legend = {
        "Math": "Math growth, numeracy, STAAR math, Algebra readiness, student practice.",
        "MTSS": "Intervention, Tier 2/Tier 3, progress monitoring, campus variation.",
        "SPED/ELL": "Special education, multilingual learners, subgroup access, differentiated supports.",
        "CCMR": "College, career, and military readiness; CTE; pathways; career awareness.",
        "Teacher Capacity": "Professional learning, coaching, PLCs, teacher burden, implementation support.",
        "Curriculum / HQIM": "Curriculum adoption, HQIM, Bluebonnet, Eureka, implementation fidelity.",
        "Funding / Grants": "Title funds, grants, LASSO, implementation funding alignment.",
        "Existing Relationship": "A relationship or contract signal exists in the workbook.",
    }

    for tag, meaning in legend.items():
        st.markdown(f"{badge_html(tag)} {meaning}", unsafe_allow_html=True)


# ============================================================
# MATRIX VIEW
# ============================================================

def render_opportunity_matrix(cards):
    st.markdown('<div class="section-title">Matrix View</div>', unsafe_allow_html=True)
    st.markdown(
        '<div class="helper-text">A short in-app matrix for quick scanning. Use the Word download for the full matrix.</div>',
        unsafe_allow_html=True,
    )

    if not cards:
        st.info("No cards to show.")
        return

    df = pd.DataFrame([
        {
            "District": c.get("name", ""),
            "Tier": c.get("tier", ""),
            "Score": c.get("score", ""),
            "Why It Matters": c.get("compact", {}).get("why", ""),
            "Best Entry Point": c.get("compact", {}).get("entry", ""),
            "Best Question": (c.get("questions") or [""])[0],
            "Contacts": "; ".join(c.get("contacts", [])[:3]),
        }
        for c in cards
    ])

    st.dataframe(df, use_container_width=True, hide_index=True)

    st.download_button(
        "Download short matrix as CSV",
        data=df.to_csv(index=False).encode("utf-8"),
        file_name="district_opportunity_matrix.csv",
        mime="text/csv",
    )


# ============================================================
# APP LAYOUT
# ============================================================

st.markdown('<div class="main-title">Strategic District Field Guide</div>', unsafe_allow_html=True)
st.markdown(
    '<div class="subtitle">Quick district refresh cards for conferences, plus deeper matrix exports for follow-up planning.</div>',
    unsafe_allow_html=True,
)

with st.sidebar:
    st.header("1. Data")
    uploaded_file = st.file_uploader("Optional: upload workbook override", type=["xlsx"])
    st.caption("If no workbook is uploaded, the app will use the built-in workbook from GitHub.")

workbook_source, data_source_label = get_workbook_source(uploaded_file)

if workbook_source is None:
    tab_field, tab_howto, tab_matrix, tab_diag = st.tabs(
        ["Field Guide", "How-To Guide", "Matrix View", "Workbook Diagnostics"]
    )
    with tab_field:
        st.error(f"No workbook found. Add {BUILT_IN_WORKBOOK} to the GitHub repo or upload a workbook in the sidebar.")
    with tab_howto:
        render_how_to(data_source_label)
    with tab_matrix:
        st.info("Workbook required to populate the matrix.")
    with tab_diag:
        st.info("Workbook required to view diagnostics.")
    st.stop()

(
    cards,
    xls,
    strategic_df,
    scorecard_df,
    basic_df,
    contacts_df,
    leadership_df,
    csi_df,
    tsi_df,
    audit_df,
    score_cols,
) = load_cards_from_workbook(workbook_source)

if not cards:
    tab_field, tab_howto, tab_matrix, tab_diag = st.tabs(
        ["Field Guide", "How-To Guide", "Matrix View", "Workbook Diagnostics"]
    )
    with tab_field:
        st.warning("No eligible district cards were generated. Review the Workbook Diagnostics tab.")
    with tab_howto:
        render_how_to(data_source_label)
    with tab_matrix:
        st.info("No eligible cards available.")
    with tab_diag:
        show_workbook_debug(
            xls,
            strategic_df,
            scorecard_df,
            basic_df,
            contacts_df,
            leadership_df,
            csi_df,
            tsi_df,
            audit_df,
            score_cols,
            data_source_label,
        )
    st.stop()

all_districts = [card.get("name", "") for card in cards]
all_tiers = sorted({card.get("tier", "") for card in cards if card.get("tier")})
all_tags = sorted({tag for card in cards for tag in card.get("tags", [])})

with st.sidebar:
    st.success(f"Using: {data_source_label}")
    st.header("2. Search")
    query = st.text_input("Search", placeholder="District, contact, signal, offering...")

    st.header("3. Filters")
    selected_tiers = st.multiselect("Tier", options=all_tiers, default=[])
    selected_tags = st.multiselect("Strategic / solution tags", options=all_tags, default=[])

    st.header("4. Shortlist")
    shortlist = st.multiselect("Shortlist districts", options=all_districts, default=[])
    show_shortlist_only = st.checkbox("Show shortlisted only", value=False)

    st.header("5. Display")
    view_mode = st.radio("View mode", options=["Quick Brief", "Full Detail"], index=0)
    st.caption("Tip: For conferences, use Quick Brief. For prep or follow-up, download the full matrix.")

filtered_cards = filter_cards(
    cards,
    query,
    selected_tiers,
    selected_tags,
    shortlist,
    show_shortlist_only,
)

tab_field, tab_howto, tab_matrix, tab_diag = st.tabs(
    ["Field Guide", "How-To Guide", "Matrix View", "Workbook Diagnostics"]
)

with tab_field:
    metric1, metric2, metric3, metric4 = st.columns(4)
    metric1.metric("📍 Cards Shown", len(filtered_cards))
    metric2.metric("✅ Eligible", len(cards))
    metric3.metric(
        "🔥 High Priority",
        sum(1 for card in filtered_cards if card.get("priority") in ["High", "Very High"]),
    )
    metric4.metric(
        "⭐ Tier 1",
        sum(1 for card in filtered_cards if card.get("tier") == "Tier 1"),
    )

    if query:
        st.markdown(
            f'<div class="metric-note">Showing results for: <strong>{safe_html(query)}</strong></div>',
            unsafe_allow_html=True,
        )

    col_dl1, col_dl2, col_dl3 = st.columns([1, 1, 1])

    with col_dl1:
        st.download_button(
            "Download brief Word guide",
            data=build_docx(filtered_cards),
            file_name="Strategic_District_Field_Guide.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            disabled=not filtered_cards,
        )

    with col_dl2:
        st.download_button(
            "Download full matrix",
            data=build_matrix_docx(filtered_cards),
            file_name="Texas_District_Strategic_Positioning_Matrix.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            disabled=not filtered_cards,
        )

    with col_dl3:
        selected_brief = st.selectbox(
            "One-district brief",
            options=[""] + [c["name"] for c in filtered_cards],
        )
        if selected_brief:
            selected_card = next(c for c in filtered_cards if c["name"] == selected_brief)
            st.download_button(
                "Download selected brief",
                data=build_docx([selected_card]),
                file_name=f"{selected_brief.replace(' ', '_')}_Brief.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                key="download_selected_brief",
            )

    st.divider()

    if not filtered_cards:
        st.info("No district cards match the current search/filter.")
    else:
        for card in filtered_cards:
            render_card(card, view_mode=view_mode)
            st.divider()

with tab_howto:
    render_how_to(data_source_label)

with tab_matrix:
    render_opportunity_matrix(filtered_cards)
    st.download_button(
        "Download full matrix as Word document",
        data=build_matrix_docx(filtered_cards),
        file_name="Texas_District_Strategic_Positioning_Matrix.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        disabled=not filtered_cards,
        key="matrix_tab_download",
    )

with tab_diag:
    show_workbook_debug(
        xls,
        strategic_df,
        scorecard_df,
        basic_df,
        contacts_df,
        leadership_df,
        csi_df,
        tsi_df,
        audit_df,
        score_cols,
        data_source_label,
    )
``
